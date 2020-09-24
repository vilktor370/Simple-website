from bs4 import BeautifulSoup
from docx import Document
import os
import requests
import webbrowser
import wget
 

    
 
 
''' 
code=''
course=''
course_code=input().upper().replace(' ','')
for i in course_code:
    if i.isdigit()==True:
        code+=i
    elif i=='X':
        code+=i
    else:
        course+=i
'''
CSE=['231','232','260','320','335','325','331']
MTH=['132','133','234','235','299','309','310','320','314']
REL=['220','101','306']

code=REL
course='REL'

path = os.path.join(os.getcwd(), course) 
 
os.mkdir(path)
 
for i in code:
    
  
 
    url = "https://msugrades.com/course/{}/{}".format(course,i)
    
    
    r = requests.get(url)
    soup=BeautifulSoup(r.text,'lxml')
    content=soup.body.find('div',id='grade-content')
    i=1
    list_length=len(content.find_all('div',class_='tab-pane fade'))
    course_title=content.find('div',id='tab0').h3.text
     
    csv =Document() 
     
    csv.add_heading(str(course_title))
    csv.add_paragraph('\n\n')
    
    
    while i<=list_length:
        
       
        div=content.find('div',id='tab'+str(i))
        data={}
        prof=div.h3.text
        prof_info=str(div.find_all('p')[1].text).replace('\n',' ')
        data[prof]=prof_info
        final=str(data).replace("{",'').replace('}','  ').replace("'",'')
        csv.add_paragraph(final)
        csv.add_paragraph('\n\n')
        i+=1
        csv_path=os.path.join(path,str(course_title)+'.docx')
        csv.save(csv_path)
  
 
       
     

 

    
    
    
 
    
 
    
    
     
 


 


 




 
 
 
 
    
 


 

